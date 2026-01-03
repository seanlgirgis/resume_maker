import yaml
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def load_config(path):
    with open(path, 'r', encoding='utf-8') as f:
        return yaml.safe_load(f)

def hex_to_rgb(hex_str):
    hex_str = hex_str.lstrip('#')
    return RGBColor(int(hex_str[0:2], 16), int(hex_str[2:4], 16), int(hex_str[4:6], 16))

def set_cell_border(cell, **kwargs):
    """
    Helper to set cell borders.
    c.f. https://stackoverflow.com/questions/33069697/how-to-setup-cell-borders-with-python-docx
    Usage: set_cell_border(cell, top={"sz": 12, "color": "FF0000", "val": "single"})
    """
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = tcPr.first_child_found_in("w:tcBorders")
    if tcBorders is None:
        tcBorders = OxmlElement('w:tcBorders')
        tcPr.append(tcBorders)

    for edge in ('left', 'top', 'right', 'bottom', 'insideH', 'insideV'):
        edge_data = kwargs.get(edge)
        if edge_data:
            tag = 'w:{}'.format(edge)
            element = tcBorders.find(qn(tag))
            if element is None:
                element = OxmlElement(tag)
                tcBorders.append(element)
            
            for key in ["sz", "val", "color", "space", "shadow"]:
                if key in edge_data:
                    element.set(qn('w:{}'.format(key)), str(edge_data[key]))

def add_shading(cell, color_hex):
    """
    Add background shading to a table cell.
    """
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), color_hex)
    tcPr.append(shd)

def create_resume(config_path, output_path):
    config = load_config(config_path)
    theme = config.get('theme', {})
    primary_color = hex_to_rgb(theme.get('primary_color', '004080'))
    
    doc = Document()
    
    # 1. Header
    header_data = config['header']
    h1 = doc.add_paragraph()
    h1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = h1.add_run(header_data['name'])
    run.bold = True
    run.font.size = Pt(24)
    run.font.color.rgb = primary_color
    run.font.name = theme.get('font_header', 'Calibri')

    contact_line = doc.add_paragraph()
    contact_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_text = f"{header_data['location']} • {header_data['phone']} • {header_data['email']}"
    run_c = contact_line.add_run(contact_text)
    run_c.font.size = Pt(10)
    
    # Links
    if 'links' in header_data:
        run_c = contact_line.add_run(" • ")
        for i, link in enumerate(header_data['links']):
            run_link = contact_line.add_run(link['text'])
            run_link.font.color.rgb = primary_color
            run_link.font.underline = True
            if i < len(header_data['links']) - 1:
                contact_line.add_run(" • ")

    doc.add_paragraph() # Spacer

    # 2. Summary (Shaded Box)
    summary_data = config['summary']
    # Create a 1x1 table for the box effect
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    
    # Add vertical accent bar logic here if needed, for now using left border
    set_cell_border(cell, left={"sz": 24, "color": theme.get('accent_color', 'E07000'), "val": "single"})
    add_shading(cell, "F8F9FA") # Light gray background
    
    # Content inside the box
    p = cell.paragraphs[0]
    runner = p.add_run(summary_data['title'])
    runner.bold = True
    p.add_run(f" {summary_data['subtitle']}\n\n")
    p.add_run(summary_data['text'])
    
    doc.add_paragraph() # Spacer

    # 3. Core Pillars
    pillars = config['core_pillars']
    h_pillars = doc.add_paragraph()
    run_p = h_pillars.add_run(pillars['title'])
    run_p.bold = True
    run_p.font.color.rgb = primary_color
    run_p.font.underline = True
    
    # 2-column table
    table_pillars = doc.add_table(rows=1, cols=2)
    table_pillars.autofit = True
    
    for idx, col_data in enumerate(pillars['columns']):
        cell = table_pillars.cell(0, idx)
        p = cell.paragraphs[0]
        run_h = p.add_run(col_data['header'])
        run_h.bold = True
        p.add_run("\n")
        
        for item in col_data['items']:
            # Simple bullet point simulation
            p.add_run(f"• {item}\n")

    doc.add_paragraph() # Spacer

    # 4. Flagship Project
    flagship = config['flagship_project']
    h_flag = doc.add_paragraph()
    run_f = h_flag.add_run(flagship['section_title'])
    run_f.bold = True
    run_f.font.color.rgb = primary_color
    run_f.font.underline = True
    
    # Boxed Project
    table_proj = doc.add_table(rows=1, cols=1)
    cell_proj = table_proj.cell(0, 0)
    set_cell_border(cell_proj, 
                    top={"sz": 6, "color": "CCCCCC", "val": "single"},
                    bottom={"sz": 6, "color": "CCCCCC", "val": "single"},
                    left={"sz": 24, "color": theme.get('accent_color', 'E07000'), "val": "single"},
                    right={"sz": 6, "color": "CCCCCC", "val": "single"})
    
    p = cell_proj.paragraphs[0]
    run_pt = p.add_run(flagship['project_title'])
    run_pt.bold = True
    p.add_run("\n\n")
    
    for highlight in flagship['highlights']:
        p.add_run(f"• {highlight}\n")
    
    p.add_run("\n")
    # Tags (naive implementation)
    tags = " | ".join(flagship['tags'])
    run_tags = p.add_run(f"[{tags}]")
    run_tags.bold = True
    run_tags.font.color.rgb = hex_to_rgb("FFFFFF")
    run_tags.font.highlight_color = 1 # Black highlight? docx is limited here, maybe just text color
    # Note: python-docx doesn't do "pills" easily without complex XML. 
    # For now, we'll just bold them or maybe use a background color if valid.

    doc.add_paragraph() # Spacer

    # 5. Experience
    exp = config['professional_experience']
    h_exp = doc.add_paragraph()
    run_e = h_exp.add_run(exp['title'])
    run_e.bold = True
    run_e.font.color.rgb = primary_color
    run_e.font.underline = True

    for job in exp['jobs']:
        p_job = doc.add_paragraph()
        r_comp = p_job.add_run(job['company'])
        r_comp.bold = True
        r_comp.font.color.rgb = primary_color
        
        p_job.add_run(f" | {job['role']}")
        
        # Right aligned date is hard in same paragraph without tabs, using string append for now
        # OR use a table for Company | Date layout
        t_job = doc.add_table(rows=1, cols=2)
        t_job.autofit = False
        t_job.columns[0].width = Inches(5.0)
        t_job.columns[1].width = Inches(1.5)
        
        c1 = t_job.cell(0, 0)
        c2 = t_job.cell(0, 1)
        
        # We replace the paragraph approach with this table approach for better layout
        c1.text = "" 
        p_c1 = c1.paragraphs[0]
        r_c1 = p_c1.add_run(job['company'])
        r_c1.bold = True
        r_c1.font.color.rgb = primary_color
        p_c1.add_run(f"\n{job['role']}")
        
        c2.text = job['dates']
        c2.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Highlights
        if 'highlights' in job:
            for highlight in job['highlights']:
                # Basic list
                doc.add_paragraph(f"• {highlight}", style='List Bullet')

    doc.save(output_path)
    print(f"Resume generated at: {output_path}")

if __name__ == "__main__":
    create_resume('resume_config.yaml', 'generated_resume.docx')
